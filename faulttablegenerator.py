import os
import csv
import docx

## This python script converts a *.csv file that contains application fault code content into
## a word table.

# A. Ensure our working directory is set to where the python file exists.
#    1) Find the path where the script is run from and set the absolute path to this.
#    2) Change directory to absolute path.
#    3) Return the working director to 'cwd' for use downstream.
dir_path = os.path.dirname(os.path.realpath(__file__))
abspath = os.path.abspath(dir_path) 
os.chdir(abspath)
cwd = os.getcwd()

# B. Pull in the CSV file's info and update list. 
#    1) Request input from user for file name.
#    2) Read in the *.csv file.
#    3) Skip the first row (i.e., headers)
#    4) Declare a faultdata list.  This will end up being a list of lists.
#    5) Loop through the csv and append each row (i.e., list) to our faultdata list.
csvFile = input("Enter name of the *.csv file, then press ENTER: ")

with open(cwd + '\\' + csvFile, 'r') as theCSVFile:
    csv_reader = csv.reader(theCSVFile)
    next(csv_reader)
    faultdata = list()
    for line in csv_reader:
        faultdata.append(line)

# C. Create a word document from faultdata list.
#    1) Define word document, theDocument.
#    2) Add a general heading to the word file.
#    3) Add a table to the word document, with 1 row and 3 columns.  The columns
#       will be a fixed size.  More rows will be appended to the table downstream.
#    4) Define the style of the table.
#    5) Define the header cells which will reside in the first row.
#    6) Loop through faultdata list and add rows that will populate the three columns
#       left-to-right with fault code, fault text, and fault remedy, respectively.
#    7) Save the document and launch it.
theDocument = docx.Document()
theDocument.add_heading('Fault Code Table', 0)
menuTable = theDocument.add_table(rows = 1, cols = 3)
menuTable.style = 'Table Grid'
header_values = menuTable.rows[0].cells
header_values[0].text = 'FAULT'
header_values[1].text = 'FAULT TEXT'
header_values[2].text = 'REMEDY'

for faultcode, faulttext, faultremedy in faultdata:
	row_values = menuTable.add_row().cells
	row_values[0].text = faultcode
	row_values[1].text = faulttext
	row_values[2].text = faultremedy
	
theDocument.save(csvFile + '.docx')

# Auto launch/open the word file
# os.system("start faultcodetable.docx")

# References:
# 1) https://github.com/RadiantCoding/Code
#    Search for "Word Doc - Simple Table"
